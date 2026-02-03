#!/usr/bin/env python3
"""
Script to create CNI Fleet Management Automation document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level):
    """Add a heading with consistent formatting"""
    heading = doc.add_heading(text, level=level)
    return heading

def add_table(doc, headers, rows):
    """Add a formatted table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        set_cell_shading(header_cells[i], '1F4E79')
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_data):
            row.cells[i].text = str(cell_text)

    return table

def create_document():
    """Create the main document"""
    doc = Document()

    # Title
    title = doc.add_heading('FLEET MANAGEMENT AUTOMATION', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Requirements Summary & System Design Discussion')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle2 = doc.add_paragraph('C&I Leasing (CNI) × Liberty Assured')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Document info table
    doc.add_paragraph()
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('Meeting Date:', 'January 15, 2026'),
        ('Prepared By:', 'Kpongette Inyang, Technical Lead, Liberty Assured'),
        ('Document Date:', 'January 16, 2026 (Updated)')
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        for paragraph in info_table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    doc.add_page_break()

    # Section 1: Executive Summary
    add_heading(doc, '1. EXECUTIVE SUMMARY', 1)

    doc.add_paragraph(
        'Thank you for the comprehensive discussion on January 15, 2026. This document synthesizes '
        'CNI Leasing\'s fleet management automation requirements from the meeting, organized into '
        'eight core functional areas, and addresses critical design questions around user touchpoints '
        'and workflow integration.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Core Objective: ').bold = True
    p.add_run(
        'Replace manual, Excel-based processes with an integrated digital system that seamlessly '
        'connects booking, maintenance, tracking, billing, supply chain, and client management — '
        'all synchronized with Sage X3 ERP and Ganoli GPS tracking.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Key Design Questions to Address:').bold = True

    bullets = [
        'How do different user types (drivers, support staff, maintenance, operations, finance, management) interact with the system?',
        'Do they need different interfaces?',
        'Should we integrate with current Excel workflows during transition or move directly to new interfaces?'
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet, style='List Bullet')

    return doc

def add_section_2(doc):
    """Section 2: Current State Assessment"""
    add_heading(doc, '2. CURRENT STATE ASSESSMENT', 1)

    add_heading(doc, '2.1 Existing Systems & Infrastructure', 2)

    systems = [
        ('Legacy FMS:', 'PHP-based fleet management system built in 2016. Supports basic vehicle registration, document tracking, and fueling approvals but lacks automation and modern integration capabilities.'),
        ('Sage X3 ERP:', 'Central financial system of record. Supports REST API integration. All invoicing, payments, POs, and financial reporting must sync here.'),
        ('Ganoli GPS Tracking:', 'Third-party GPS tracking system. Currently provides vehicle location via separate software. API integration needed to pull real-time mileage, fuel levels, location data into new system.'),
        ('Manual Processes:', 'Heavy reliance on Excel spreadsheets for booking logs, maintenance schedules, fuel tracking, invoice generation. Manual data entry into Sage. Up to 50 contract invoices generated manually each month.'),
        ('Infrastructure:', 'Microsoft Azure cloud hosting with Linux servers. Dedicated domain and IP address.'),
        ('IT Team:', 'Led by Oladejo Lasisi (Head of IT) with Nicodemus Yobo and additional capable staff ready for collaborative development and eventual handover.')
    ]

    for label, desc in systems:
        p = doc.add_paragraph()
        p.add_run(label).bold = True
        p.add_run(' ' + desc)

    add_heading(doc, '2.2 Key Pain Points', 2)

    pain_points = [
        'Booking process starts with customer phone calls, followed by manual follow-ups and payment tracking — causing inefficiencies and delays',
        'No call logging system — cannot track call volume, answered vs. missed calls, or link calls to bookings',
        'Manual invoice generation and Sage entry for both reservations and contracts (50+ contract invoices/month)',
        'Maintenance managed manually through PO/job order process — no automated alerts or historical tracking to prevent recurring repairs',
        'No real-time fleet visibility — must call drivers to get location, mileage, status updates',
        'Fuel consumption tracked manually without ability to detect fraud or anomalies',
        'Document expiries (insurance, roadworthy certificates) managed manually — risk of fines and voided coverage',
        'Cannot analyze fleet profitability per vehicle or per client due to scattered data',
        'No centralized supply chain management — parts, inventory, and procurement handled ad-hoc',
        'Outsourcing/recruitment processes (interviews, offer letters, payroll) done manually'
    ]

    for point in pain_points:
        doc.add_paragraph(point, style='List Bullet')

    return doc

def add_section_3_intro(doc):
    """Section 3: Core Functional Requirements Intro"""
    doc.add_page_break()
    add_heading(doc, '3. CORE FUNCTIONAL REQUIREMENTS', 1)

    doc.add_paragraph(
        'Based on the January 15 meeting and subsequent analysis, CNI requires automation organized '
        'into eight core functional areas:'
    )

    modules = [
        ('3.1', 'Client Management'),
        ('3.2', 'Dashboard / Business Overview Board'),
        ('3.3', 'Driver Management'),
        ('3.4', 'Fleet Management'),
        ('3.5', 'Invoice, Billing & Financial Management (Expanded)'),
        ('3.6', 'Maintenance Schedule & Management'),
        ('3.7', 'Supply Chain Management (New)'),
        ('3.8', 'Outsourcing & Recruitment Automation (Optional)')
    ]

    for num, name in modules:
        p = doc.add_paragraph()
        p.add_run(f'{num} ').bold = True
        p.add_run(name)

    return doc

def add_section_3_1(doc):
    """Section 3.1: Client Management"""
    add_heading(doc, '3.1 Client Management', 2)

    doc.add_paragraph(
        'Objective: Centralized customer relationship management with complete interaction history, '
        'call tracking, and contract oversight.'
    )

    # Customer Database
    add_heading(doc, 'Customer Database', 3)
    features = [
        ('Complete customer master data', 'Name, company, contact details, ID documents'),
        ('Booking history per customer', 'All past and current bookings'),
        ('Payment history and credit terms', 'Payment records and agreed terms'),
        ('Customer segmentation', 'Walk-in, corporate, contract, VIP'),
        ('Sage X3 Integration', 'Sync with Sage X3 customer accounts')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # Call Tracking
    add_heading(doc, 'Call Tracking & Logging', 3)

    p = doc.add_paragraph()
    p.add_run('Requirement from meeting: ').bold = True
    p.add_run('"Everything originates from a customer calling. We want to have visibility around calls — '
              'how many calls received, was this call successful, who picked it, was it bounced, '
              'nobody picked at all" — Mobolaji Johnson')

    doc.add_paragraph()

    features = [
        ('Call logging system', 'Track all customer fleet requests'),
        ('Call details recording', 'Date/time, duration, caller ID, staff who answered'),
        ('Call status tracking', 'Answered vs. missed calls'),
        ('Call-to-booking linking', 'Link calls to resulting bookings/reservations'),
        ('Call analytics dashboard', 'Volume, response times, conversion rates'),
        ('VoIP/SIP integration', 'Manual logging initially, VoIP integration in future phase')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # Contract Management
    add_heading(doc, 'Contract Management', 3)
    features = [
        ('Contract vehicle database', 'Clients with monthly vehicle rentals'),
        ('Contract terms storage', 'Rates, duration, renewal dates'),
        ('Renewal alerts', 'Automated alerts at 90, 60, 30 days before expiry'),
        ('Amendment workflow', 'Process for contract changes'),
        ('SLA tracking', 'Compliance reporting')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # Client Self-Service Portal
    add_heading(doc, 'Client Self-Service Portal (Optional)', 3)
    features = [
        ('Vehicle/booking view', 'Clients view their vehicles, bookings, invoices'),
        ('Online requests', 'Request vehicles or services online'),
        ('Online payments', 'Make payments through portal'),
        ('Document downloads', 'Download invoices and reports')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    return doc


def add_section_3_2(doc):
    """Section 3.2: Dashboard / Business Overview Board"""
    doc.add_page_break()
    add_heading(doc, '3.2 Dashboard / Business Overview Board', 2)

    doc.add_paragraph(
        'Objective: Real-time operational visibility and KPI tracking for management decision-making '
        'and profitability analysis.'
    )

    # Executive Dashboard
    add_heading(doc, 'Executive Dashboard', 3)
    features = [
        ('High-level KPIs', 'Daily/monthly revenue, fleet utilization %, active bookings, maintenance alerts'),
        ('Financial overview', 'Revenue trends, outstanding invoices, payment collection rate'),
        ('Fleet status summary', 'Vehicles active, idle, in-maintenance, dormant'),
        ('Critical alerts', 'Overdue maintenance, expiring documents, missed calls, payment delays')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # Live Fleet Monitoring
    add_heading(doc, 'Live Fleet Monitoring', 3)

    p = doc.add_paragraph()
    p.add_run('Requirement from meeting: ').bold = True
    p.add_run('"We want to be able to track all these fleets. A system where we can look at and say, '
              'okay, this fleet is in so place or perhaps it\'s in a workshop, it\'s off. Why is it off? '
              'What is going on" — Mobolaji Johnson')

    doc.add_paragraph()

    features = [
        ('Live map display', 'Show all vehicle locations (Ganoli GPS integration)'),
        ('Color-coded status', 'Active (green), idle (yellow), in-workshop (orange), offline (red), dormant (gray)'),
        ('Vehicle details on click', 'Current assignment, driver, client, last update'),
        ('Filter & search', 'Filter by status, client, driver, location')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # Profitability Analysis
    add_heading(doc, 'Profitability Analysis', 3)

    p = doc.add_paragraph()
    p.add_run('Requirement from meeting: ').bold = True
    p.add_run('"Profitability report is why we need this development. We want all this data to come in '
              'through a source so we can have everything in one bucket" — Mobolaji Johnson')

    doc.add_paragraph()

    features = [
        ('Per-vehicle profitability', 'Revenue vs. all costs (fuel, maintenance, insurance, depreciation)'),
        ('Per-client profitability', 'Profitability per client/contract'),
        ('Cost breakdown', 'Detailed cost categorization and trending'),
        ('Budget vs. actual', 'Compare planned vs. actual spending'),
        ('Sage X3 sync', 'Pull financial data from Sage for accurate reporting')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # Custom Reports
    add_heading(doc, 'Custom Reports', 3)
    features = [
        ('Report builder', 'Create custom queries and reports'),
        ('Scheduled reports', 'Automate daily/weekly/monthly report generation'),
        ('Export options', 'Export to Excel, PDF, CSV'),
        ('Email distribution', 'Auto-send reports to stakeholders')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    return doc


def add_section_3_3(doc):
    """Section 3.3: Driver Management"""
    doc.add_page_break()
    add_heading(doc, '3.3 Driver Management', 2)

    doc.add_paragraph(
        'Objective: Complete driver lifecycle from onboarding through performance tracking, '
        'compliance management, and payroll integration.'
    )

    # Driver Database
    add_heading(doc, 'Driver Database', 3)
    features = [
        ('Personal details', 'Employment history, contact information'),
        ('License tracking', 'Driver\'s license details with expiry tracking'),
        ('Certifications', 'Training certifications and medical fitness certificates'),
        ('Document expiry alerts', 'Automated alerts for license, medical, certifications'),
        ('Auto-deactivation', 'Automatically deactivate driver if license expires')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # Vehicle-Driver Assignment
    add_heading(doc, 'Vehicle-Driver Assignment', 3)
    features = [
        ('Assignment management', 'Assign drivers to vehicles for bookings'),
        ('Assignment history', 'Track driver-vehicle assignment history'),
        ('Rotation management', 'Manage driver rotation schedules'),
        ('Availability calendar', 'View and manage driver scheduling')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # Performance Tracking
    add_heading(doc, 'Performance Tracking', 3)
    features = [
        ('Performance scoring', 'Score based on on-time delivery, fuel efficiency, incidents, customer ratings'),
        ('Driving behavior', 'Track harsh braking, speeding (if GPS supports)'),
        ('Performance bonuses', 'Link performance to compensation/bonuses')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # Driver Mobile Interface
    add_heading(doc, 'Driver Mobile Interface', 3)
    features = [
        ('Trip logging', 'Log trip start/end with mileage and fuel level'),
        ('Vehicle condition photos', 'Capture photos at handover/return'),
        ('Maintenance requests', 'Submit requests with photos'),
        ('Incident reporting', 'Report incidents or damages'),
        ('Fuel authorization', 'Request fuel authorization'),
        ('Offline capability', 'Sync when network available'),
        ('Digital signature', 'Capture e-signatures'),
        ('Simple UI', 'Intuitive interface for low-tech literacy users')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # Payroll Integration
    add_heading(doc, 'Payroll Integration', 3)
    features = [
        ('Salary information', 'Store driver salary/wage data'),
        ('Bonuses & deductions', 'Track performance bonuses and deductions'),
        ('Sage X3 integration', 'Sync with Sage for payroll processing')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    return doc


def add_section_3_4(doc):
    """Section 3.4: Fleet Management"""
    doc.add_page_break()
    add_heading(doc, '3.4 Fleet Management', 2)

    doc.add_paragraph(
        'Objective: Complete vehicle lifecycle management with real-time tracking, reservation system, '
        'document compliance, and fuel management.'
    )

    # Vehicle Master Database
    add_heading(doc, 'Vehicle Master Database', 3)
    features = [
        ('Vehicle details', 'Make, model, year, VIN, registration number'),
        ('Acquisition details', 'Purchase date, cost, supplier'),
        ('Valuation tracking', 'Current valuation and depreciation'),
        ('Specifications', 'Seating capacity, cargo space, features'),
        ('Photos & documents', 'Store vehicle images and documents'),
        ('Status management', 'Available, Booked, Active, In-Maintenance, In-Workshop, Dormant')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # GPS Tracking
    add_heading(doc, 'Real-Time GPS Tracking Integration (Ganoli)', 3)
    features = [
        ('Real-time location', 'Latitude, longitude, street address'),
        ('Speed & heading', 'Current speed and direction'),
        ('Odometer reading', 'Mileage for maintenance scheduling and fuel reconciliation'),
        ('Fuel level', 'If tracker hardware supports'),
        ('Ignition status', 'On/off status'),
        ('Battery monitoring', 'Battery voltage tracking'),
        ('Update timestamp', 'Last update time and GPS signal quality'),
        ('Route playback', 'Historical route viewing'),
        ('Alerts', 'Vehicle offline >2 hours, unauthorized movement, geofence breaches')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # Booking & Reservation
    add_heading(doc, 'Booking & Reservation Management', 3)

    p = doc.add_paragraph()
    p.add_run('Requirement from meeting: ').bold = True
    p.add_run('"Customer requests fleet, we show them availability, they make payment online, '
              'payment confirmed, invoice generated automatically, ticket issued for driver and security, '
              'vehicle released" — Mobolaji Johnson')

    doc.add_paragraph()
    doc.add_paragraph('Customer-Facing Portal:', style='List Bullet')

    features = [
        ('Vehicle browsing', 'Browse available vehicles with photos and specs'),
        ('Availability calendar', 'Check availability with conflict detection'),
        ('Rate calculator', 'Calculate rates based on duration'),
        ('Online payment', 'Integrated with Wema/GT Bank/VFD'),
        ('Instant confirmation', 'Immediate booking confirmation')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    doc.add_paragraph('Staff Interface:', style='List Bullet')

    features = [
        ('Phone/walk-in bookings', 'Create bookings for phone/walk-in customers'),
        ('Status tracking', 'Pending, Confirmed, Active, Completed, Cancelled'),
        ('Driver assignment', 'Assign driver to booking'),
        ('Digital ticket/QR code', 'Generate for vehicle release')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    doc.add_paragraph('Handover Process:', style='List Bullet')

    features = [
        ('Pre-trip checklist', 'Inspection, mileage, fuel level, photos'),
        ('Return process', 'Post-trip inspection, mileage verification, damage documentation')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    return doc


def add_section_3_4_continued(doc):
    """Section 3.4 continued: Document Compliance and Fuel Management"""

    # Document Compliance
    add_heading(doc, 'Document Compliance Management', 3)

    doc.add_paragraph('Centralized document repository per vehicle:')
    docs = ['Insurance certificate', 'Roadworthy certificate', 'Vehicle registration',
            'Inspection reports', 'Purchase documents']
    for d in docs:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('Automated expiry tracking with multi-level alerts:')

    features = [
        ('90 days before', 'Informational alert'),
        ('30 days before', 'Warning alert (daily reminders)'),
        ('7 days before', 'Critical alert (multiple daily reminders)'),
        ('On expiry', 'Vehicle marked unavailable'),
        ('Compliance dashboard', 'Color-coded: green, yellow, orange, red'),
        ('Renewal workflow', 'Track renewals and vendors'),
        ('Version history', 'Document version tracking')
    ]
    add_table(doc, ['Alert Level', 'Action'], features)

    doc.add_paragraph()

    # Fuel Management
    add_heading(doc, 'Fuel Management', 3)

    p = doc.add_paragraph()
    p.add_run('Requirement from meeting: ').bold = True
    p.add_run('"We want to see the fuel level, trace the distance with whatever fueling level. '
              'We can\'t justify the cost — fueling yesterday, coming back again today" — Mobolaji Johnson')

    doc.add_paragraph()
    doc.add_paragraph('Fuel Request Workflow:', style='List Bullet')

    features = [
        ('Driver request', 'Driver initiates via mobile app'),
        ('System display', 'Shows last fuel date and current mileage'),
        ('Approval workflow', 'Based on amount thresholds'),
        ('Fuel code/PO', 'Generate authorized fuel code'),
        ('Receipt submission', 'Driver submits receipt with photo')
    ]
    add_table(doc, ['Step', 'Description'], features)

    doc.add_paragraph()
    doc.add_paragraph('Consumption Tracking:', style='List Bullet')

    features = [
        ('Fuel efficiency', 'Calculate km per liter'),
        ('Expected vs. actual', 'Compare consumption'),
        ('Cost per kilometer', 'Track fuel cost efficiency'),
        ('Budget vs. actual', 'Compare planned vs. actual spending')
    ]
    add_table(doc, ['Metric', 'Description'], features)

    doc.add_paragraph()
    doc.add_paragraph('Fraud Detection Alerts:', style='List Bullet')

    features = [
        ('No mileage increase', 'Fuel purchased but no mileage change'),
        ('Excessive consumption', '>30% above vehicle average'),
        ('Frequent fueling', '>3x per week'),
        ('After-hours fueling', 'Unusual timing'),
        ('Price discrepancies', 'Unusual pricing')
    ]
    add_table(doc, ['Alert Type', 'Trigger Condition'], features)

    return doc


def add_section_3_5(doc):
    """Section 3.5: Invoice, Billing & Financial Management (Expanded)"""
    doc.add_page_break()
    add_heading(doc, '3.5 Invoice, Billing & Financial Management (EXPANDED)', 2)

    doc.add_paragraph(
        'Objective: Fully automated invoice generation, payment processing, and accounts receivable/payable '
        'management with seamless Sage X3 integration.'
    )

    # 3.5.1 Invoice Types & Generation
    add_heading(doc, '3.5.1 Invoice Types & Generation', 3)

    # Reservation Invoices
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Reservation/Booking Invoices').bold = True

    features = [
        ('Auto-generation trigger', 'Automatic invoice upon booking confirmation'),
        ('Invoice number format', 'Configurable format (e.g., CNI-RES-2026-00001)'),
        ('Invoice details', 'Booking reference, vehicle, dates, rate, total, VAT, payment terms'),
        ('Rate calculation', 'Auto-calculate based on duration, vehicle type, client tier, seasonal pricing'),
        ('Additional charges', 'Extra mileage, fuel charges, late return fees, damage charges, cleaning fees'),
        ('Discounts', 'Percentage or fixed discounts, promotional codes, loyalty discounts'),
        ('PDF generation', 'CNI-branded PDF with logo, terms & conditions, bank details'),
        ('Multi-currency', 'Support for USD, EUR, GBP alongside NGN'),
        ('Auto-email', 'Email invoice to customer with PDF attachment'),
        ('Sage X3 sync', 'Real-time posting to Sage as sales invoice')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Contract Monthly Invoices').bold = True

    features = [
        ('Automated scheduling', 'Auto-generate on billing day (configurable: 1st, 15th, last day)'),
        ('Batch processing', 'Process all contract clients simultaneously'),
        ('Contract reference', 'Link to master contract with terms'),
        ('Period specification', 'Billing period (e.g., "February 2026")'),
        ('Vehicle listing', 'All vehicles on contract with individual rates'),
        ('Variable charges', 'Additional mileage, fuel, driver overtime, ad-hoc services'),
        ('Fixed charges', 'Monthly rental, insurance, maintenance package'),
        ('Pro-rata calculation', 'Handle mid-month additions/removals of vehicles'),
        ('Escalation clauses', 'Auto-apply annual rate increases per contract terms'),
        ('Consolidated invoicing', 'Option to consolidate multiple contracts per client'),
        ('Split invoicing', 'Option to split by department/cost center for large clients'),
        ('Auto-distribution', 'Email to designated finance contacts per client'),
        ('Approval workflow', 'Optional internal approval before sending (for large invoices)')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Ad-Hoc/Miscellaneous Invoices').bold = True

    features = [
        ('Manual creation', 'Create invoices for non-standard services'),
        ('Service categories', 'Consulting, training, special transport, event services, penalties'),
        ('Line item flexibility', 'Add unlimited line items with descriptions, quantities, rates'),
        ('Template library', 'Save and reuse invoice templates for common ad-hoc services'),
        ('Attachment support', 'Attach supporting documents (photos, reports, agreements)')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Credit Notes').bold = True

    features = [
        ('Credit note generation', 'Issue credit notes against original invoices'),
        ('Reason codes', 'Overcharge, Service not rendered, Discount adjustment, Goodwill, Error correction'),
        ('Partial credits', 'Credit specific line items or partial amounts'),
        ('Full reversal', 'Complete invoice reversal'),
        ('Approval workflow', 'Require manager approval for credits above threshold'),
        ('Auto-application', 'Option to auto-apply credit to next invoice or refund'),
        ('Sage X3 sync', 'Post credit note to Sage')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Debit Notes').bold = True

    features = [
        ('Debit note generation', 'Issue for additional charges post-invoice'),
        ('Use cases', 'Damage discovered after return, additional mileage, penalty charges'),
        ('Link to original', 'Reference original invoice/booking'),
        ('Supporting evidence', 'Attach photos, inspection reports, GPS data'),
        ('Client notification', 'Auto-notify client with explanation')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Proforma Invoices').bold = True

    features = [
        ('Quotation generation', 'Create proforma for client approval before service'),
        ('Validity period', 'Set expiry date for proforma (e.g., valid for 14 days)'),
        ('Conversion to invoice', 'One-click conversion to final invoice'),
        ('Version tracking', 'Track revisions if client requests changes'),
        ('Approval status', 'Track client approval/rejection'),
        ('PO linkage', 'Link received PO to proforma before conversion')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    return doc


def add_section_3_5_continued(doc):
    """Section 3.5 continued: PO Management and Payment Processing"""

    # 3.5.2 Purchase Order Management
    add_heading(doc, '3.5.2 Purchase Order (PO) Management', 3)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Client PO Handling').bold = True

    features = [
        ('PO receipt tracking', 'Record client POs received'),
        ('PO validation', 'Match PO to proforma/quote, validate amounts'),
        ('PO attachment', 'Store scanned PO documents'),
        ('PO-invoice linking', 'Mandatory PO reference on invoices for corporate clients'),
        ('PO consumption tracking', 'Track partial invoicing against blanket POs'),
        ('PO expiry alerts', 'Alert when PO value/date nearing exhaustion')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Supplier PO Generation').bold = True

    features = [
        ('PO creation', 'Generate POs for maintenance, fuel, parts, services'),
        ('PO number format', 'Configurable format (e.g., CNI-PO-2026-00001)'),
        ('Approval workflow', 'Multi-level approval based on amount thresholds'),
        ('Budget checking', 'Validate against department/project budgets'),
        ('Vendor selection', 'Select from approved vendor list'),
        ('PO status tracking', 'Draft, Pending Approval, Approved, Sent, Partially Received, Completed, Cancelled'),
        ('Sage X3 sync', 'Post approved POs to Sage')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # 3.5.3 Payment Processing
    add_heading(doc, '3.5.3 Payment Processing & Collections', 3)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Payment Gateway Integration').bold = True

    features = [
        ('Wema Bank (Cash Connect)', 'Primary online payment gateway'),
        ('VFD Microfinance Bank', 'Alternative payment option'),
        ('GT Bank', 'Corporate banking integration'),
        ('Card payments', 'Accept Visa, Mastercard, Verve'),
        ('Bank transfer', 'Accept and auto-reconcile bank transfers'),
        ('USSD payments', 'Mobile payment support'),
        ('Payment confirmation', 'Real-time payment confirmation and receipt')
    ]
    add_table(doc, ['Gateway/Method', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Payment Recording').bold = True

    features = [
        ('Multiple payment types', 'Cash, check, bank transfer, card, mobile money'),
        ('Partial payments', 'Accept and track partial payments'),
        ('Overpayments', 'Handle overpayments (credit balance or refund)'),
        ('Payment date', 'Record actual payment date vs. receipt date'),
        ('Reference numbers', 'Bank reference, check number, transaction ID'),
        ('Payment proof', 'Attach payment evidence (bank statement, receipt)')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Payment Allocation').bold = True

    features = [
        ('Auto-allocation', 'Automatically allocate payments to oldest invoices (FIFO)'),
        ('Manual allocation', 'Option to manually allocate to specific invoices'),
        ('Split payments', 'Allocate single payment across multiple invoices'),
        ('Unapplied payments', 'Track unallocated payments'),
        ('Allocation reversal', 'Correct allocation errors'),
        ('Sage X3 sync', 'Post payments and allocations to Sage')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Payment Receipts').bold = True

    features = [
        ('Auto-generation', 'Generate receipt upon payment recording'),
        ('Receipt format', 'CNI-branded PDF receipt'),
        ('Receipt number', 'Unique receipt numbering'),
        ('Email delivery', 'Auto-email receipt to payer'),
        ('Receipt history', 'Store all receipts for audit')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Refund Processing').bold = True

    features = [
        ('Refund request', 'Initiate refund against credit note or overpayment'),
        ('Approval workflow', 'Require approval for refunds'),
        ('Refund methods', 'Bank transfer, original payment method'),
        ('Refund tracking', 'Track refund status through completion'),
        ('Sage X3 sync', 'Post refund to Sage')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    return doc


def add_section_3_5_ar_ap(doc):
    """Section 3.5 continued: AR and AP Management"""

    # 3.5.4 Accounts Receivable Management
    add_heading(doc, '3.5.4 Accounts Receivable Management', 3)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Invoice Aging').bold = True

    features = [
        ('Aging buckets', 'Current, 1-30 days, 31-60 days, 61-90 days, 90+ days'),
        ('Aging reports', 'Detailed aging by client, by invoice, summary'),
        ('Aging dashboard', 'Visual representation of AR health'),
        ('Trend analysis', 'Track aging trends over time')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Automated Dunning').bold = True

    features = [
        ('Reminder schedule', 'Configurable reminder schedule (e.g., 7, 14, 30, 60 days overdue)'),
        ('Email templates', 'Customizable reminder templates per aging level'),
        ('Escalation path', 'Auto-escalate to manager/collections after threshold'),
        ('Client-specific rules', 'Different dunning rules per client/segment'),
        ('Reminder history', 'Log all reminders sent'),
        ('Pause capability', 'Option to pause reminders for disputed invoices')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Collections Management').bold = True

    features = [
        ('Collections queue', 'Prioritized list of overdue accounts'),
        ('Collection actions', 'Log calls, emails, visits, promises to pay'),
        ('Promise to pay', 'Record and track payment promises'),
        ('Collection status', 'Track collection stage (reminder, call, legal)'),
        ('Write-off process', 'Workflow for bad debt write-off with approvals'),
        ('Collection agency', 'Track accounts sent to external collection')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    return doc



def add_section_3_5_ap_tax(doc):
    """Section 3.5 continued: AP and Tax Management"""

    # 3.5.5 Accounts Payable Management
    add_heading(doc, '3.5.5 Accounts Payable Management', 3)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Vendor Invoice Processing').bold = True

    features = [
        ('Invoice receipt', 'Record vendor invoices with scan/upload'),
        ('PO matching', '3-way match: PO, goods receipt, invoice'),
        ('Discrepancy handling', 'Flag and workflow for mismatches'),
        ('Invoice approval', 'Multi-level approval workflow'),
        ('Due date tracking', 'Track payment due dates'),
        ('Early payment discount', 'Flag early payment discounts available')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Payment Scheduling').bold = True

    features = [
        ('Payment calendar', 'View upcoming payment obligations'),
        ('Cash flow forecast', 'Forecast cash requirements'),
        ('Payment prioritization', 'Prioritize by due date, discount, vendor importance'),
        ('Batch payments', 'Group payments for efficiency'),
        ('Hold payments', 'Option to hold payments pending resolution')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Payment Execution').bold = True

    features = [
        ('Payment methods', 'Bank transfer, check, cash'),
        ('Bank file generation', 'Generate bank payment files for bulk transfers'),
        ('Payment approval', 'Dual approval for payments above threshold'),
        ('Payment confirmation', 'Record payment confirmation and reference'),
        ('Sage X3 sync', 'Post vendor payments to Sage')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # 3.5.6 Tax Management
    add_heading(doc, '3.5.6 Tax Management', 3)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('VAT/Tax Handling').bold = True

    features = [
        ('Tax configuration', 'Configure VAT rates (7.5% standard, exempt, zero-rated)'),
        ('Auto-calculation', 'Automatically calculate tax on invoices'),
        ('Tax-inclusive pricing', 'Option for tax-inclusive or tax-exclusive pricing'),
        ('Tax exemption', 'Handle tax-exempt clients/services'),
        ('Tax breakdown', 'Show tax breakdown on invoices')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Withholding Tax').bold = True

    features = [
        ('WHT rates', 'Configure WHT rates by service type'),
        ('WHT calculation', 'Calculate WHT on applicable invoices'),
        ('WHT certificate', 'Generate/receive WHT certificates'),
        ('WHT reconciliation', 'Reconcile WHT credits')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Tax Reporting').bold = True

    features = [
        ('VAT returns', 'Generate data for VAT returns'),
        ('WHT reports', 'Withholding tax reports'),
        ('Tax audit trail', 'Complete audit trail for tax transactions'),
        ('FIRS compliance', 'Reports formatted for Nigerian tax authority')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    return doc


def add_section_3_5_reporting_sage(doc):
    """Section 3.5 continued: Financial Reporting and Sage X3 Integration"""

    # 3.5.7 Financial Reporting
    add_heading(doc, '3.5.7 Financial Reporting & Analytics', 3)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Revenue Reports').bold = True

    features = [
        ('Revenue by period', 'Daily, weekly, monthly, quarterly, annual'),
        ('Revenue by client', 'Revenue breakdown by client/client segment'),
        ('Revenue by vehicle', 'Revenue per vehicle for profitability'),
        ('Revenue by service', 'Breakdown by reservation, contract, ad-hoc'),
        ('Revenue trends', 'Trend analysis and forecasting')
    ]
    add_table(doc, ['Report', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Receivables Reports').bold = True

    features = [
        ('Outstanding invoices', 'All unpaid invoices with aging'),
        ('Collection forecast', 'Expected collection timeline'),
        ('DSO calculation', 'Days Sales Outstanding metrics'),
        ('Bad debt analysis', 'Analysis of write-offs and provisions')
    ]
    add_table(doc, ['Report', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Payables Reports').bold = True

    features = [
        ('Outstanding payables', 'All unpaid vendor invoices'),
        ('Payment forecast', 'Upcoming payment obligations'),
        ('DPO calculation', 'Days Payable Outstanding metrics'),
        ('Vendor spend analysis', 'Spending by vendor/category')
    ]
    add_table(doc, ['Report', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Profitability Reports').bold = True

    features = [
        ('Gross margin', 'Revenue minus direct costs'),
        ('Vehicle profitability', 'Profit/loss per vehicle (revenue - fuel - maintenance - depreciation)'),
        ('Client profitability', 'Profit/loss per client'),
        ('Contract profitability', 'Margin analysis per contract'),
        ('Cost center reporting', 'P&L by department/cost center')
    ]
    add_table(doc, ['Report', 'Description'], features)

    doc.add_paragraph()

    # 3.5.8 Sage X3 Integration
    add_heading(doc, '3.5.8 Sage X3 ERP Integration (Expanded)', 3)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Outbound to Sage X3').bold = True

    features = [
        ('Sales invoices', 'Post all invoices to Sage sales ledger'),
        ('Credit notes', 'Post credit notes to Sage'),
        ('Customer receipts', 'Post customer payments'),
        ('Purchase orders', 'Post approved POs to Sage purchasing'),
        ('Vendor payments', 'Post supplier payments'),
        ('Journal entries', 'Post adjusting entries as needed')
    ]
    add_table(doc, ['Transaction Type', 'Integration'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Inbound from Sage X3').bold = True

    features = [
        ('Chart of accounts', 'Sync GL account structure'),
        ('Customer master', 'Sync customer data'),
        ('Vendor master', 'Sync vendor/supplier data'),
        ('Bank transactions', 'Import bank postings for reconciliation'),
        ('Period status', 'Check if periods are open/closed')
    ]
    add_table(doc, ['Data Type', 'Integration'], features)

    return doc



def add_section_3_6(doc):
    """Section 3.6: Maintenance Schedule & Management"""
    doc.add_page_break()
    add_heading(doc, '3.6 Maintenance Schedule & Management', 2)

    doc.add_paragraph(
        'Objective: Proactive maintenance scheduling, complete maintenance history, '
        'workshop management, and cost tracking to minimize downtime and extend vehicle life.'
    )

    # Automated Maintenance Scheduling
    add_heading(doc, 'Automated Maintenance Scheduling', 3)

    p = doc.add_paragraph()
    p.add_run('Requirement from meeting: ').bold = True
    p.add_run('"We want to be able to track that vehicle X, every time you see an issue with the vehicle, '
              'you know maintenance schedule says maybe 1000 km servicing, you see it as issue. '
              'Something is wrong" — Mobolaji Johnson')

    doc.add_paragraph()

    features = [
        ('Service schedules', 'Define maintenance by mileage intervals (e.g., every 5,000 km) or time (e.g., every 3 months)'),
        ('GPS mileage sync', 'Pull real-time mileage from Ganoli to trigger schedules'),
        ('Upcoming alerts', 'Alert at configurable thresholds (e.g., 500 km before due)'),
        ('Overdue alerts', 'Escalating alerts when maintenance is past due'),
        ('Auto-scheduling', 'Automatically create work orders when thresholds reached'),
        ('Recurring tasks', 'Define recurring maintenance tasks by vehicle type'),
        ('Seasonal maintenance', 'Schedule seasonal checks (e.g., AC before summer)')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # Maintenance Request Workflow
    add_heading(doc, 'Maintenance Request Workflow', 3)

    features = [
        ('Driver-initiated requests', 'Drivers submit maintenance needs via mobile app with photos'),
        ('Manager review', 'Operations manager reviews and approves/modifies'),
        ('Priority assignment', 'Set priority: Emergency, High, Medium, Low'),
        ('Work order creation', 'Convert request to work order with job details'),
        ('Workshop assignment', 'Assign to internal workshop or external vendor'),
        ('Parts requisition', 'Link required parts to work order'),
        ('Status tracking', 'Track: Pending, In Progress, Awaiting Parts, Completed'),
        ('Completion sign-off', 'Digital sign-off on completed work'),
        ('Quality inspection', 'Post-maintenance inspection checklist')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # Maintenance History
    add_heading(doc, 'Complete Maintenance History', 3)

    features = [
        ('Full history', 'Complete maintenance record per vehicle'),
        ('Cost tracking', 'Track parts cost, labor cost, external service cost'),
        ('Pattern detection', 'Identify recurring issues for same vehicle'),
        ('Warranty tracking', 'Track parts/work under warranty'),
        ('Document storage', 'Store invoices, receipts, photos per maintenance'),
        ('Vendor performance', 'Track vendor quality and turnaround time')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # Workshop Management
    add_heading(doc, 'Workshop Management', 3)

    features = [
        ('Bay management', 'Manage workshop bays/capacity'),
        ('Technician assignment', 'Assign technicians to jobs'),
        ('Time tracking', 'Track time spent on each job'),
        ('Parts consumption', 'Track parts used from inventory'),
        ('Job costing', 'Calculate total cost per job'),
        ('Capacity planning', 'Schedule based on available capacity'),
        ('External vendor tracking', 'Manage external workshop relationships')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # Maintenance Analytics
    add_heading(doc, 'Maintenance Analytics', 3)

    features = [
        ('Cost per vehicle', 'Total maintenance cost per vehicle over time'),
        ('Cost per km', 'Maintenance cost per kilometer traveled'),
        ('MTBF', 'Mean Time Between Failures by vehicle/type'),
        ('Downtime analysis', 'Days out of service for maintenance'),
        ('Budget vs. actual', 'Maintenance budget tracking'),
        ('Predictive alerts', 'Flag vehicles with increasing maintenance costs')
    ]
    add_table(doc, ['Metric', 'Description'], features)

    return doc



def add_section_3_7(doc):
    """Section 3.7: Supply Chain Management (NEW)"""
    doc.add_page_break()
    add_heading(doc, '3.7 Supply Chain Management (NEW MODULE)', 2)

    doc.add_paragraph(
        'Objective: End-to-end supply chain visibility from procurement through inventory management, '
        'vendor relationships, and warehouse operations to support fleet maintenance and operations.'
    )

    # 3.7.1 Inventory Management
    add_heading(doc, '3.7.1 Inventory Management', 3)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Parts & Spares Inventory').bold = True

    features = [
        ('Item master', 'Comprehensive parts catalog with descriptions, specifications, photos'),
        ('SKU management', 'Unique SKU for each part with barcode/QR support'),
        ('Categorization', 'Hierarchical categories (Engine, Electrical, Body, Tires, Fluids, etc.)'),
        ('Vehicle compatibility', 'Link parts to compatible vehicle makes/models'),
        ('Unit of measure', 'Support multiple UoMs (piece, liter, set, pair)'),
        ('Location tracking', 'Track location within warehouse (aisle, shelf, bin)'),
        ('Batch/serial tracking', 'Track batches for fluids, serial numbers for major parts')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Stock Levels & Movements').bold = True

    features = [
        ('Real-time quantities', 'Current stock levels per item per location'),
        ('Minimum stock levels', 'Define reorder points per item'),
        ('Maximum stock levels', 'Define maximum inventory to prevent overstocking'),
        ('Stock movements', 'Track receipts, issues, transfers, adjustments'),
        ('Goods receipt', 'Record incoming inventory against POs'),
        ('Goods issue', 'Issue parts to work orders/maintenance'),
        ('Stock transfer', 'Transfer between locations/warehouses'),
        ('Stock adjustment', 'Adjust for variances with reason codes'),
        ('Cycle counting', 'Schedule and record cycle counts')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Inventory Alerts').bold = True

    features = [
        ('Low stock alert', 'Alert when stock falls below reorder point'),
        ('Overstock alert', 'Alert when stock exceeds maximum'),
        ('Slow-moving items', 'Flag items with no movement for X days'),
        ('Expiry alerts', 'Alert for items approaching expiry (oils, fluids)'),
        ('Negative stock', 'Prevent or alert on negative stock conditions')
    ]
    add_table(doc, ['Alert Type', 'Description'], features)

    doc.add_paragraph()

    # 3.7.2 Procurement Management
    add_heading(doc, '3.7.2 Procurement Management', 3)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Purchase Requisition').bold = True

    features = [
        ('Requisition creation', 'Create purchase requests for parts/services'),
        ('Requester details', 'Track who requested and department'),
        ('Urgency levels', 'Standard, Urgent, Emergency'),
        ('Approval workflow', 'Multi-level approval based on value'),
        ('Conversion to PO', 'Convert approved requisitions to POs'),
        ('Consolidation', 'Combine multiple requisitions into single PO')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Supplier Sourcing').bold = True

    features = [
        ('RFQ generation', 'Create Request for Quotation for competitive pricing'),
        ('Vendor comparison', 'Compare quotations from multiple vendors'),
        ('Price history', 'Track historical prices for items'),
        ('Preferred vendors', 'Maintain preferred vendor per item category'),
        ('Contract pricing', 'Use contracted pricing for recurring purchases')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Order Tracking').bold = True

    features = [
        ('PO status', 'Track: Created, Sent, Acknowledged, Shipped, Delivered'),
        ('Expected delivery', 'Track expected delivery dates'),
        ('Partial receipts', 'Handle partial deliveries'),
        ('Delivery discrepancies', 'Track quantity/quality discrepancies'),
        ('Returns processing', 'Process returns to vendors')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    return doc


def add_section_3_7_continued(doc):
    """Section 3.7 continued: Vendor, Fuel, and Tire Management"""

    # 3.7.3 Vendor/Supplier Management
    add_heading(doc, '3.7.3 Vendor/Supplier Management', 3)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Vendor Master').bold = True

    features = [
        ('Vendor details', 'Company name, contacts, addresses, bank details'),
        ('Product categories', 'What categories vendor supplies'),
        ('Payment terms', 'Agreed payment terms'),
        ('Tax information', 'VAT registration, WHT requirements'),
        ('Certification tracking', 'Track vendor certifications/licenses'),
        ('Active/inactive status', 'Manage vendor lifecycle')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Vendor Performance').bold = True

    features = [
        ('Delivery performance', 'On-time delivery rate'),
        ('Quality rating', 'Track defects/returns per vendor'),
        ('Price competitiveness', 'Compare pricing vs. market'),
        ('Response time', 'Track quotation response times'),
        ('Overall scoring', 'Composite vendor score'),
        ('Performance reviews', 'Schedule periodic vendor reviews')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # 3.7.4 Fuel Supply Chain
    add_heading(doc, '3.7.4 Fuel Supply Chain', 3)

    features = [
        ('Fuel procurement', 'Track bulk fuel purchases'),
        ('Depot management', 'Manage fuel depot inventory (if applicable)'),
        ('Fuel card management', 'Track fuel cards issued to drivers'),
        ('Station agreements', 'Manage agreements with fuel stations'),
        ('Fuel pricing', 'Track fuel prices over time'),
        ('Vendor reconciliation', 'Reconcile fuel station invoices with consumption data')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # 3.7.5 Tire Management
    add_heading(doc, '3.7.5 Tire Management', 3)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Tire Inventory').bold = True

    features = [
        ('Tire specifications', 'Size, brand, type, speed rating'),
        ('New tire stock', 'Track new tires in inventory'),
        ('Retread stock', 'Track retreaded tires'),
        ('Scrap tires', 'Track tires removed from service')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Tire Lifecycle Tracking').bold = True

    features = [
        ('Tire-to-vehicle assignment', 'Track which tire on which vehicle position'),
        ('Mileage tracking', 'Track km per tire'),
        ('Rotation history', 'Record tire rotations'),
        ('Repair history', 'Track punctures and repairs'),
        ('Retread cycles', 'Track number of retreads'),
        ('Retirement criteria', 'Flag tires for replacement')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    return doc



def add_section_3_7_final(doc):
    """Section 3.7 final: Asset Management, Warehouse, Analytics"""

    # 3.7.6 Asset Management
    add_heading(doc, '3.7.6 Asset Management', 3)

    features = [
        ('Fixed asset register', 'Track all fixed assets (vehicles, equipment, tools)'),
        ('Asset tagging', 'Barcode/QR code for asset identification'),
        ('Depreciation tracking', 'Calculate and track depreciation'),
        ('Asset lifecycle', 'Track acquisition, transfers, disposal'),
        ('Asset valuation', 'Current book value per asset'),
        ('Insurance tracking', 'Link assets to insurance policies'),
        ('Sage X3 sync', 'Sync asset data with Sage fixed assets module')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # 3.7.7 Warehouse Management
    add_heading(doc, '3.7.7 Warehouse Management', 3)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Warehouse Setup').bold = True

    features = [
        ('Multiple warehouses', 'Support multiple storage locations'),
        ('Location structure', 'Define zones, aisles, racks, shelves, bins'),
        ('Location types', 'Receiving, storage, staging, shipping'),
        ('Capacity planning', 'Track storage capacity utilization')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Warehouse Operations').bold = True

    features = [
        ('Receiving', 'Process inbound deliveries with inspection'),
        ('Putaway', 'Assign storage locations for received items'),
        ('Picking', 'Pick items for work orders/maintenance'),
        ('Packing', 'Prepare items for internal/external transfer'),
        ('Shipping', 'Record outbound shipments'),
        ('Returns handling', 'Process returns to warehouse')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # 3.7.8 Supply Chain Analytics
    add_heading(doc, '3.7.8 Supply Chain Analytics & Reporting', 3)

    features = [
        ('Inventory value', 'Total inventory value by category/location'),
        ('Inventory turnover', 'Turnover rate analysis'),
        ('Stock aging', 'Age of inventory items'),
        ('Procurement spend', 'Spending by vendor/category/period'),
        ('Lead time analysis', 'Vendor lead time tracking'),
        ('Cost trends', 'Price trends for key items'),
        ('Demand forecasting', 'Predict future requirements based on history'),
        ('ABC analysis', 'Classify items by value/movement')
    ]
    add_table(doc, ['Report', 'Description'], features)

    doc.add_paragraph()

    # 3.7.9 Supply Chain Integration
    add_heading(doc, '3.7.9 Supply Chain Integration', 3)

    features = [
        ('Maintenance integration', 'Link parts consumption to work orders'),
        ('Financial integration', 'Inventory valuation to Sage X3'),
        ('Vendor payment integration', 'Payables from goods receipt'),
        ('Fleet integration', 'Fuel and tire data to vehicle records'),
        ('Mobile integration', 'Warehouse operations via mobile devices')
    ]
    add_table(doc, ['Integration', 'Description'], features)

    return doc


def add_section_3_8(doc):
    """Section 3.8: Outsourcing & Recruitment Automation"""
    doc.add_page_break()
    add_heading(doc, '3.8 Outsourcing & Recruitment Automation (Optional)', 2)

    doc.add_paragraph(
        'Objective: Streamline recruitment and onboarding process for contract staff '
        '(drivers, security, support) including interviews, assessments, and payroll setup.'
    )

    # Job Posting & Applications
    add_heading(doc, 'Job Posting & Applications', 3)

    features = [
        ('Job requisition', 'Create job requests with approval workflow'),
        ('Job posting', 'Post openings (internal and external channels)'),
        ('Application tracking', 'Collect and track applications'),
        ('Resume parsing', 'Extract key information from resumes'),
        ('Candidate database', 'Maintain candidate pool for future positions'),
        ('Communication', 'Automated emails to candidates at each stage')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # Interview & Assessment
    add_heading(doc, 'Interview & Assessment', 3)

    features = [
        ('Interview scheduling', 'Schedule interviews with calendar integration'),
        ('Assessment tests', 'Online driving tests, skill assessments'),
        ('Interview scoring', 'Standardized scoring rubrics'),
        ('Background checks', 'Track status of background verification'),
        ('Document verification', 'Verify licenses, certifications, references'),
        ('Hiring decision', 'Collaborative decision workflow')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    doc.add_paragraph()

    # Onboarding
    add_heading(doc, 'Onboarding', 3)

    features = [
        ('Offer letter generation', 'Auto-generate offer letters'),
        ('Document collection', 'Checklist for required documents'),
        ('Contract signing', 'Electronic signature for contracts'),
        ('Asset issuance', 'Track uniforms, ID cards, fuel cards issued'),
        ('Training tracking', 'Schedule and track mandatory training'),
        ('Payroll setup', 'Create employee record and sync to Sage'),
        ('System access', 'Create accounts in relevant systems')
    ]
    add_table(doc, ['Feature', 'Description'], features)

    return doc



def add_section_4(doc):
    """Section 4: User Types & Touchpoints"""
    doc.add_page_break()
    add_heading(doc, '4. USER TYPES & TOUCHPOINTS', 1)

    doc.add_paragraph(
        'The system must support different user types with appropriate interfaces and permissions:'
    )

    user_types = [
        ('Drivers', 'Mobile app for trip logging, fuel requests, maintenance reporting, document submission'),
        ('Support Staff', 'Web interface for call logging, booking creation, customer service'),
        ('Operations Team', 'Dashboard for fleet monitoring, driver assignment, vehicle status'),
        ('Maintenance Team', 'Work order management, parts requisition, job completion'),
        ('Finance Team', 'Invoice management, payment processing, AR/AP, Sage integration'),
        ('Procurement Team', 'PO creation, vendor management, inventory management'),
        ('Management', 'Executive dashboard, reports, approvals, analytics'),
        ('Clients', 'Self-service portal for bookings, invoices, payments (optional)'),
        ('IT Admin', 'System configuration, user management, integrations')
    ]
    add_table(doc, ['User Type', 'Primary Interface & Functions'], user_types)

    return doc


def add_section_5(doc):
    """Section 5: Workflow Integration Strategy"""
    add_heading(doc, '5. WORKFLOW INTEGRATION STRATEGY', 1)

    doc.add_paragraph(
        'Recommended approach to transition from current manual/Excel processes:'
    )

    strategies = [
        ('Phase 1: Parallel Running', 'Run new system alongside Excel for 1-2 months, compare outputs'),
        ('Phase 2: Primary System', 'New system becomes primary, Excel as backup only'),
        ('Phase 3: Full Transition', 'Complete migration, Excel discontinued for operational use'),
        ('Data Import', 'Initial data migration from Excel and legacy systems'),
        ('Training', 'Comprehensive user training before each phase'),
        ('Change Management', 'Regular feedback sessions and system adjustments')
    ]
    add_table(doc, ['Phase/Strategy', 'Description'], strategies)

    return doc


def add_section_6(doc):
    """Section 6: Key Integration Requirements"""
    add_heading(doc, '6. KEY INTEGRATION REQUIREMENTS', 1)

    doc.add_paragraph('Critical system integrations required:')

    integrations = [
        ('Sage X3 ERP', 'REST API', 'Invoices, payments, POs, GL, customers, vendors'),
        ('Ganoli GPS', 'API', 'Real-time vehicle location, mileage, fuel level, speed'),
        ('Payment Gateways', 'API', 'Wema Bank, VFD, GT Bank for online payments'),
        ('Email/SMS', 'SMTP/API', 'Notifications, alerts, invoice delivery, reminders'),
        ('Mobile Apps', 'REST API', 'Driver app (Android/iOS) for field operations'),
        ('Document Storage', 'Azure Blob', 'Store documents, photos, attachments')
    ]
    add_table(doc, ['System', 'Integration Type', 'Data/Purpose'], integrations)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Integration Contacts:').bold = True

    contacts = [
        'Sage X3: Oladejo Lasisi (Head of IT)',
        'Ganoli GPS: To be provided by CNI',
        'Payment Gateways: Finance team contacts'
    ]
    for c in contacts:
        doc.add_paragraph(c, style='List Bullet')

    return doc


def add_section_7(doc):
    """Section 7: Data Migration Plan"""
    add_heading(doc, '7. DATA MIGRATION PLAN', 1)

    doc.add_paragraph('Data to be migrated from existing systems:')

    migrations = [
        ('Vehicle Master', 'Legacy FMS, Excel', 'High', 'All active vehicles with history'),
        ('Driver Records', 'Legacy FMS, Excel', 'High', 'All active drivers'),
        ('Customer Master', 'Sage X3, Excel', 'High', 'All active customers'),
        ('Active Contracts', 'Excel', 'High', 'All current contracts'),
        ('Open Invoices', 'Sage X3', 'High', 'All unpaid invoices'),
        ('Maintenance History', 'Excel, Legacy FMS', 'Medium', 'Last 2-3 years'),
        ('Fuel Records', 'Excel', 'Medium', 'Last 12 months for baseline'),
        ('Document Library', 'File shares', 'Medium', 'Current valid documents'),
        ('Vendor Master', 'Sage X3', 'Medium', 'Active vendors'),
        ('Parts Catalog', 'Excel', 'Medium', 'If available')
    ]
    add_table(doc, ['Data Type', 'Source', 'Priority', 'Scope'], migrations)

    return doc


def add_section_8(doc):
    """Section 8: Project Approach & Methodology"""
    doc.add_page_break()
    add_heading(doc, '8. PROJECT APPROACH & METHODOLOGY', 1)

    doc.add_paragraph(
        'Liberty Assured proposes an Agile-Scrum methodology with phased delivery:'
    )

    # Proposed Phases
    add_heading(doc, 'Proposed Implementation Phases', 2)

    phases = [
        ('Phase 1', 'Booking & Reservations, Client Management, Call Logging', 'Revenue-generating, immediate business impact'),
        ('Phase 2', 'Invoice & Billing, Payment Processing, Sage X3 Integration', 'Automates 50+ manual invoices, cash flow impact'),
        ('Phase 3', 'Fleet Management, GPS Integration, Fuel Management', 'Real-time visibility, fraud detection'),
        ('Phase 4', 'Maintenance Scheduling, Work Order Workflow, Driver Management', 'Cost control, downtime reduction'),
        ('Phase 5', 'Dashboard & Analytics, Profitability Reports, Custom Reports', 'Management decision-making'),
        ('Phase 6', 'Supply Chain, Inventory, Procurement', 'Operational efficiency'),
        ('Phase 7', 'Document Management, Compliance Alerts', 'Risk mitigation'),
        ('Phase 8', 'Outsourcing/Recruitment Module (Optional)', 'Separate business line')
    ]
    add_table(doc, ['Phase', 'Modules', 'Rationale'], phases)

    doc.add_paragraph()

    # Sprint Structure
    add_heading(doc, 'Sprint Structure', 2)

    items = [
        '2-week sprints with defined deliverables',
        'Sprint planning at start of each sprint',
        'Daily standups for progress tracking',
        'Sprint demos for stakeholder feedback',
        'Sprint retrospectives for continuous improvement'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    return doc



def add_section_9(doc):
    """Section 9: Information Needed from CNI"""
    add_heading(doc, '9. INFORMATION NEEDED FROM CNI', 1)

    doc.add_paragraph('To proceed with detailed design and development, Liberty Assured requires:')

    # Operational Details
    add_heading(doc, 'Operational Details', 2)
    items = [
        'Complete vehicle list with specifications and current status',
        'Driver roster with license and certification details',
        'Current contract list with terms and billing schedules',
        'Sample invoices (reservation, contract, ad-hoc)',
        'Current booking/reservation process documentation',
        'Maintenance schedule templates and history samples',
        'Fuel management current process and vendor agreements'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # Technical Details
    add_heading(doc, 'Technical Details', 2)
    items = [
        'Sage X3 API documentation and credentials (sandbox)',
        'Ganoli GPS API documentation and access',
        'Azure infrastructure details and access',
        'Legacy FMS database schema (if migration needed)',
        'Current Excel templates used for operations'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # Business Priorities
    add_heading(doc, 'Business Priorities', 2)
    items = [
        'Prioritized list of pain points to address first',
        'Key stakeholders for each module',
        'Decision-makers for approvals during development',
        'Preferred go-live timeline'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    return doc


def add_section_10(doc):
    """Section 10: Legal & Contractual Requirements"""
    add_heading(doc, '10. LEGAL & CONTRACTUAL REQUIREMENTS', 1)

    doc.add_paragraph('Standard agreements required before project commencement:')

    agreements = [
        ('Non-Disclosure Agreement (NDA)', 'Protect confidential business and technical information'),
        ('Data Protection Agreement', 'NDPR compliance for handling personal data'),
        ('Master Service Agreement', 'Define scope, deliverables, timelines, payment terms'),
        ('SLA Agreement', 'Define support levels post-deployment'),
        ('IP Assignment', 'Clarify ownership of developed software')
    ]
    add_table(doc, ['Agreement', 'Purpose'], agreements)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Cost Transparency: ').bold = True
    p.add_run('Detailed cost breakdown will be provided in a separate commercial proposal '
              'following agreement on scope and phasing.')

    return doc


def add_section_11(doc):
    """Section 11: Why Liberty Assured"""
    doc.add_page_break()
    add_heading(doc, '11. WHY LIBERTY ASSURED', 1)

    doc.add_paragraph(
        'Liberty Assured brings unique value to this engagement:'
    )

    strengths = [
        ('Fleet Management Expertise', 'Deep experience in fleet automation, GPS integration, and maintenance systems'),
        ('ERP Integration', 'Proven track record with Sage X3 and other ERP integrations'),
        ('Nigerian Market Knowledge', 'Understanding of local business practices, regulations, and payment systems'),
        ('Agile Delivery', 'Iterative approach with regular deliverables and stakeholder involvement'),
        ('Technical Excellence', 'Modern technology stack with scalable, secure architecture'),
        ('Partnership Approach', 'Collaborative development with knowledge transfer to CNI IT team'),
        ('Post-Deployment Support', 'Comprehensive support and maintenance options')
    ]
    add_table(doc, ['Strength', 'Description'], strengths)

    return doc


def add_section_12(doc):
    """Section 12: Next Steps & Timeline"""
    add_heading(doc, '12. NEXT STEPS & TIMELINE', 1)

    doc.add_paragraph('Proposed immediate next steps:')

    steps = [
        ('Week 1-2', 'CNI provides requested information and documentation'),
        ('Week 2-3', 'Liberty Assured prepares detailed technical proposal and cost estimate'),
        ('Week 3', 'Proposal presentation and Q&A session'),
        ('Week 4', 'Contract negotiation and signing'),
        ('Week 5', 'Project kickoff and Phase 1 sprint planning'),
        ('Week 6+', 'Development begins with 2-week sprint cycles')
    ]
    add_table(doc, ['Timeline', 'Activity'], steps)

    return doc


def add_section_13(doc):
    """Section 13: Conclusion"""
    add_heading(doc, '13. CONCLUSION', 1)

    doc.add_paragraph(
        'This document outlines a comprehensive Fleet Management Automation solution designed to address '
        'CNI Leasing\'s operational challenges. The proposed system will:'
    )

    benefits = [
        'Eliminate manual processes and reduce errors',
        'Provide real-time visibility into fleet operations',
        'Automate invoice generation and payment processing',
        'Enable proactive maintenance scheduling',
        'Detect and prevent fuel fraud',
        'Ensure document compliance with automated alerts',
        'Deliver actionable profitability insights',
        'Streamline supply chain and inventory management',
        'Integrate seamlessly with Sage X3 and Ganoli GPS'
    ]
    for b in benefits:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        'Liberty Assured is committed to delivering a solution that transforms CNI\'s fleet operations '
        'and provides a foundation for continued growth and efficiency.'
    )

    return doc


def add_section_14(doc):
    """Section 14: Contact Information"""
    add_heading(doc, '14. CONTACT INFORMATION', 1)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Liberty Assured').bold = True

    doc.add_paragraph('Kpongette Inyang, Technical Lead')
    doc.add_paragraph('Email: kpongette@libertyassured.com')
    doc.add_paragraph('Phone: [To be provided]')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('C&I Leasing (CNI)').bold = True

    doc.add_paragraph('Mobolaji Johnson, [Title]')
    doc.add_paragraph('Oladejo Lasisi, Head of IT')
    doc.add_paragraph('Email: [To be provided]')

    return doc


# Main execution
if __name__ == '__main__':
    print("Creating CNI Fleet Management document...")

    # Create document with initial sections
    doc = create_document()

    # Add all sections
    add_section_2(doc)
    add_section_3_intro(doc)
    add_section_3_1(doc)
    add_section_3_2(doc)
    add_section_3_3(doc)
    add_section_3_4(doc)
    add_section_3_4_continued(doc)
    add_section_3_5(doc)
    add_section_3_5_continued(doc)
    add_section_3_5_ar_ap(doc)
    add_section_3_5_ap_tax(doc)
    add_section_3_5_reporting_sage(doc)
    add_section_3_6(doc)
    add_section_3_7(doc)
    add_section_3_7_continued(doc)
    add_section_3_7_final(doc)
    add_section_3_8(doc)
    add_section_4(doc)
    add_section_5(doc)
    add_section_6(doc)
    add_section_7(doc)
    add_section_8(doc)
    add_section_9(doc)
    add_section_10(doc)
    add_section_11(doc)
    add_section_12(doc)
    add_section_13(doc)
    add_section_14(doc)

    # Save document
    doc.save('CNI UPDATED v2.docx')
    print("Document saved as 'CNI UPDATED v2.docx'")
    print("Document creation complete!")

